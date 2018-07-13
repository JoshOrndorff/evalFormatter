#! /usr/bin/env python3
#TODO can't get python-docx to install on python3

# Imports to work with the Microsoft OpenXML format
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Imports to load the template file
from os.path import dirname, isfile
from os import sep

class CTYEvaluation(object):

  '''
    Represents a CTY evaluation in both format and content. Format is handled by
    a docx.Document member, and content is passed in with calls to add_paragraph.
  '''
  def __init__(self, fname, lname, cname, date, course, instructor, site, ta, completion):
    
    # Validate required arguments (note cname is not required (can be None))
    x = 0
    for arg in [fname, lname, date, course, instructor, site, ta, completion]:
      x+=1
      if arg is None:
        print(x)
        raise ValueError("Required argument: {} was None.".format(arg))
    
    # Reference only the arguments that will be used outside of the constructor
    self.fname = fname
    self.lname = lname
    self.cname = fname if cname is None else cname
    
    # Create the document
    self.document = Document()
    
    # Set margins
    section = self.document.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title line at the very top
    title = self.document.add_paragraph()
    title.style.font.name = "Times New Roman" # All subsequent text is same font
    title.paragraph_format.space_after = Pt(11)
    title.paragraph_format.alignment = 1 # 1 means center
    title.add_run("CTY SUMMER PROGRAM FINAL EVALUATION").bold = True
        
    # Populate heading information including names, date, site, and course
    heading = self.document.add_paragraph()
    heading.paragraph_format.space_after = Pt(11)
    heading.paragraph_format.line_spacing = 1
    heading.add_run("Student: " + fname + ' ' + lname + '\t')
    heading.add_run("Date: " + date + '\n')
    heading.add_run("Course: " + course + '\t')
    heading.add_run("Instructor: " + instructor + '\n')
    heading.add_run("Site: " + site + '\t')
    heading.add_run("Teaching Assistant: " + ta)

    #TODO I guess there is better API support for tab stops in python-docx now
    # https://python-docx.readthedocs.io/en/latest/dev/analysis/features/text/tab-stops.html
    # Insert a tab stop in the middle of the page for proper heading format.
    # Simplified code from: github.com/python-openxml/python-docx/issues/206
    pTabs = OxmlElement('w:tabs')
    heading.paragraph_format.element.get_or_add_pPr().append(pTabs)
    tab_n = OxmlElement('w:tab')
    tab_n.set(qn('w:val'), "left")
    tab_n.set(qn('w:pos'), '5000')
    pTabs.append(tab_n)
    
    # Setup the intro paragraph
    introText = "Congratulations, " + self.cname + ', '
    introText += "on " + completion + ' ' + course + '. '
    introText += "Please see the enclosed course description for more detailed "
    introText += "information on the course."
    intro = self.document.add_paragraph(introText)
    intro.paragraph_format.space_after = Pt(11)
    
    intro.paragraph_format.line_spacing = 1
    # Keep track of the number of body paragraphs in the document
    self.numParagraphs = 0
    
        
  def add_paragraph(self, title, text = ""):
    ''' Add a body paragraph to the document '''
    
    if self.numParagraphs > 5:
      raise Warning("Adding sixth body paragraph. Standard CTY evaluation only"\
                     + " contains five.")
    
    paragraph = self.document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(11)
    paragraph.add_run(title + '\n').bold = True
    paragraph.add_run(text)
    
    self.numParagraphs += 1
    
  def add_notes(self, notes):
    '''
      Add a bulleted list of notes to facilitate human editing.
      Notes can be either a plain python list or an xml Element Tree.
    '''
    
    xmlStructure = type(notes) is not list
    
    for note in notes:
      if xmlStructure:
        note = note.text
      self.document.add_paragraph(note, style='List Bullet')
      
  def add_signature(self):
    ''' Adds the standard Instructor's signature line to the document '''
    signature = self.document.add_paragraph("Instructor's Signature: ")
    signature.add_run('_' * 60)

  def save(self, suffix = ""):
    '''
      Saves document to  CTY standard filename, last_first.docx, and appends
      an optional suffix.
    '''
        
    filename = self.lname + '_' + self.fname + suffix + '.docx'
      
    self.document.save(filename)
  
  def file_exists(self, suffix = ""):
    '''
      Returns boolean whether a file with the CTY standard filename,
      last_first.docx (with optional suffix esits in the pwd.
    '''
    return isfile(self.lname + '_' + self.fname + suffix + ".docx")
    
    

